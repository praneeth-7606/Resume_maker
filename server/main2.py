






# from fastapi import FastAPI, File, UploadFile, HTTPException
# from fastapi.middleware.cors import CORSMiddleware
# from fastapi.responses import FileResponse
# import pandas as pd
# import re
# import json
# from pathlib import Path
# import shutil
# import fitz
# from docx import Document
# import logging
# from mistralai import Mistral
# import os
# from datetime import datetime

# # Set up logging
# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger(__name__)

# app = FastAPI()

# # Configure CORS
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["http://localhost:3000"],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# # Create directories if they don't exist
# UPLOAD_DIR = Path("uploads")
# OUTPUT_DIR = Path("generated_documents")
# UPLOAD_DIR.mkdir(exist_ok=True)
# OUTPUT_DIR.mkdir(exist_ok=True)

# # Mistral API configuration
# MISTRAL_API_KEY = "Vl1QhEbiSbwGhfiXdKxnyz5CYqL1FgC9"
# MISTRAL_MODEL = "mistral-large-latest"

# # Template instructions
# COVER_LETTER_TEMPLATE = """
# **COVER LETTER:**
# - Start with a professional greeting.
# - Introduce yourself and mention the position applied for.
# - Highlight key relevant skills and achievements.
# - Explain why you're interested in the role and how it aligns with the company's goals.
# - End with a call to action and gratitude.
# **TONE & FORMATTING:**
# - Keep it formal, professional, and concise.
# - Ensure proper spacing and clear paragraph structure.
# """

# RESUME_TEMPLATE = """
# **PROFESSIONAL SUMMARY:**
# - Highlight key technical and behavioral competencies with scores above 4.
# - Include top 3 certifications to showcase qualifications.
# **EDUCATION:**
# - List degrees, institutions, and graduation years in reverse chronological order.
# **SKILLS:**
# - Clearly categorized technical, analytical, and soft skills.
# **PROJECT DETAILS:**
# Format projects in tabular structure:
# **<Project Name>**
# | Description | Overview of purpose and outcomes. |
# | Role | Role played in the project (e.g., Developer, Team Lead). |
# | Technology | Technologies and tools used. |
# | Contributions | Specific contributions and responsibilities. |
# **TONE & FORMATTING:**
# - Maintain a formal, professional, and structured format.
# """

# class FileProcessingError(Exception):
#     def __init__(self, message, details):
#         self.message = message
#         self.details = details
#         super().__init__(self.message)

# # [Previous helper functions remain the same up to generate_document]

# def generate_document(data, document_type="cover letter"):
#     try:
#         client = Mistral(api_key=MISTRAL_API_KEY)
        
#         # Select appropriate template based on document type
#         template = RESUME_TEMPLATE if document_type.lower() == "resume" else COVER_LETTER_TEMPLATE
        
#         prompt = f"""
#         {template}
#         Here is the extracted data:
#         {json.dumps(data, indent=2)}
#         Generate a {document_type} based on the above data.
#         """
        
#         response = client.chat.complete(
#             model=MISTRAL_MODEL,
#             messages=[{"role": "user", "content": prompt}]
#         )
        
#         return response.choices[0].message.content
#     except Exception as e:
#         logger.error(f"Error generating document: {str(e)}")
#         raise HTTPException(status_code=500, detail=f"Error generating document: {str(e)}")

# @app.post("/generate-document/{document_type}")
# async def generate_document_endpoint(document_type: str, data: dict):
#     try:
#         # Generate the document content
#         generated_content = generate_document(data, document_type)
        
#         # Create a unique filename
#         timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
#         filename = f"Generated_{document_type.replace(' ', '_')}_{timestamp}"
        
#         # Save as both DOCX and PDF
#         docx_path = OUTPUT_DIR / f"{filename}.docx"
#         pdf_path = OUTPUT_DIR / f"{filename}.pdf"
        
#         # Save as DOCX
#         doc = Document()
#         doc.add_paragraph(generated_content)
#         doc.save(docx_path)
        
#         # Convert to PDF using PyMuPDF
#         pdf_doc = fitz.open()
#         page = pdf_doc.new_page()
#         page.insert_text((50, 50), generated_content)
#         pdf_doc.save(pdf_path)
#         pdf_doc.close()
        
#         # Create a zip file containing both formats
#         zip_path = OUTPUT_DIR / f"{filename}.zip"
#         with zipfile.ZipFile(zip_path, 'w') as zipf:
#             zipf.write(docx_path, arcname=f"{filename}.docx")
#             zipf.write(pdf_path, arcname=f"{filename}.pdf")
        
#         # Clean up individual files
#         docx_path.unlink()
#         pdf_path.unlink()
        
#         return FileResponse(
#             path=zip_path,
#             filename=f"{filename}.zip",
#             media_type="application/zip"
#         )
    
#     except Exception as e:
#         logger.error(f"Error in document generation endpoint: {str(e)}")
#         raise HTTPException(status_code=500, detail=str(e))

# # Add a new endpoint for batch document generation


# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger(__name__)

# app = FastAPI()

# # Configure CORS
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["http://localhost:3000"],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# # Create directories if they don't exist
# UPLOAD_DIR = Path("uploads")
# OUTPUT_DIR = Path("generated_documents")
# UPLOAD_DIR.mkdir(exist_ok=True)
# OUTPUT_DIR.mkdir(exist_ok=True)

# # Mistral API configuration
# MISTRAL_API_KEY = "Vl1QhEbiSbwGhfiXdKxnyz5CYqL1FgC9"
# MISTRAL_MODEL = "mistral-large-latest"

# # Template instructions
# COVER_LETTER_TEMPLATE = """
# ---
# *COVER LETTER:*  
# - Start with a *professional greeting* addressing the hiring manager or company.  
# - Introduce yourself with a *brief background* and mention the *position* you're applying for.  
# - Highlight *key skills and experiences* relevant to the job, focusing on achievements and specific contributions.  
# - Explain *why you're interested* in the position and how your goals align with the company's mission.  
# - Conclude with a *call to action*, expressing eagerness for an interview and thanking them for their time.  
# ---
# *TONE & FORMATTING:*  
# - Maintain a *formal* and *professional tone* throughout.  
# - Keep it *concise* (preferably within one page) and *focused on relevant experiences*.  
# - Ensure *proper spacing* and *clear paragraph structure* for easy readability.
# ---
# """



# class FileProcessingError(Exception):
#     def __init__(self, message, details):
#         self.message = message
#         self.details = details
#         super().__init__(self.message)

# # app = FastAPI()
# # app.add_middleware(
# #     CORSMiddleware,
# #     allow_origins=["http://localhost:3000"],  # React app URL
# #     allow_credentials=True,
# #     allow_methods=["*"],
# #     allow_headers=["*"],
# # )

# # Create uploads directory if it doesn't exist
# # UPLOAD_DIR = Path("uploads")
# # UPLOAD_DIR.mkdir(exist_ok=True)

# def extract_docx(file_path):
#     doc = Document(file_path)
#     structured_text = {"paragraphs": [], "tables": []}
    
#     for para in doc.paragraphs:
#         structured_text["paragraphs"].append(para.text.strip())
    
#     for table in doc.tables:
#         table_data = []
#         for row in table.rows:
#             row_data = [cell.text.strip() for cell in row.cells]
#             table_data.append(row_data)
#         structured_text["tables"].append(table_data)
    
#     return structured_text

# def extract_pdf(file_path):
#     doc = fitz.open(file_path)
#     structured_text = {"pages": []}
    
#     for page_num in range(len(doc)):
#         page = doc[page_num]
#         text = page.get_text("text")
#         structured_text["pages"].append({
#             "page_number": page_num + 1,
#             "content": text.strip()
#         })
    
#     return structured_text

# def extract_name_and_role_from_docx(file_path):
#     doc_data = extract_docx(file_path)
#     name, role = None, None

#     for para in doc_data["paragraphs"]:
#         if not name and any(word in para.lower() for word in ["name", "full name"]):
#             name = para.split(":")[-1].strip()
#         if not role and any(word in para.lower() for word in ["role", "position", "job title"]):
#             role = para.split(":")[-1].strip()

#     return name, role

# def extract_name_and_role_from_pdf(file_path):
#     pdf_data = extract_pdf(file_path)
#     name, role = None, None

#     for page in pdf_data["pages"]:
#         content = page["content"]
#         if not name and "name" in content.lower():
#             name_line = [line for line in content.split("\n") if "name" in line.lower()][0]
#             name = name_line.split(":")[-1].strip()
#         if not role and any(word in content.lower() for word in ["role", "position", "job title"]):
#             role_line = [line for line in content.split("\n") if any(word in line.lower() for word in ["role", "position", "job title"])][0]
#             role = role_line.split(":")[-1].strip()

#     return name, role

# def filter_skill_matrix(skill_json_path, extracted_name, extracted_role):
#     with open(skill_json_path, 'r', encoding='utf-8') as file:
#         skill_data = json.load(file)

#     # Filter by role
#     filtered_by_role = [
#         entry for entry in skill_data
#         if extracted_role and extracted_role.lower() in entry.get("ID", "").lower()
#         or extracted_role.lower() in entry.get("Sheet Name", "").lower()
#     ]

#     # Further filter by name
#     filtered_by_name = []
#     for entry in filtered_by_role:
#         first_name = entry.get("First Name", "")
#         last_name = entry.get("Last Name", "")
        
#         if isinstance(first_name, str) and isinstance(last_name, str):
#             full_name = f"{first_name} {last_name}".strip().lower()
#             if extracted_name and extracted_name.lower() in full_name:
#                 filtered_by_name.append(entry)

#     return filtered_by_name

# def combine_resume_and_skills(resume_data, skill_matrix_data):
#     combined_data = {
#         "resume": resume_data,
#         "filtered_skill_matrix": skill_matrix_data
#     }
#     return combined_data
# def extract_clean_skill_matrix(file_path: Path) -> list:
#     try:
#         xls = pd.ExcelFile(file_path)
#         structured_data = []
#         role_counter = 1

#         percentage_pattern = re.compile(r"\d+%")

#         for sheet in xls.sheet_names:
#             try:
#                 df = xls.parse(sheet)
#                 df = df.dropna(how='all').fillna("")

#                 columns = list(df.columns)
#                 if len(columns) < 4:
#                     logger.warning(f"Sheet {sheet} has insufficient columns")
#                     continue

#                 first_name_col = columns[0]
#                 last_name_col = "Unnamed: 1"
#                 experience_col = columns[2]
#                 expertise_col = columns[3]

#                 categories = {
#                     "Salesforce Technical Competencies and External Systems Integration": [],
#                     "Behavioral & Leadership Competencies and Certifications": []
#                 }

#                 category_flag = None
#                 for col in columns[4:]:
#                     if isinstance(col, str):
#                         if "%" in col or "Current Capability Score" in col or "Expertise(Years)" in col:
#                             continue
#                         if "Salesforce Technical Competencies" in col or "External Systems Integration" in col:
#                             category_flag = "Salesforce Technical Competencies and External Systems Integration"
#                         elif "Behavioral & Leadership Competencies" in col or "SF Certification" in col:
#                             category_flag = "Behavioral & Leadership Competencies and Certifications"
#                         elif category_flag:
#                             categories[category_flag].append(col)

#                 for _, row in df.iterrows():
#                     first_name = str(row[first_name_col]).strip()
#                     last_name = str(row[last_name_col]).strip()
                    
#                     if first_name and last_name and first_name.lower() != "nan" and last_name.lower() != "nan":
#                         entry = {
#                             "ID": f"Role_{role_counter}",
#                             "Sheet Name": sheet,
#                             "First Name": first_name,
#                             "Last Name": last_name,
#                             "Experience": str(row[experience_col]),
#                             "Expertise": str(row[expertise_col]),
#                             "Salesforce Technical Competencies and External Systems Integration": {},
#                             "Behavioral & Leadership Competencies and Certifications": {}
#                         }

#                         # Process technical competencies
#                         for skill in categories["Salesforce Technical Competencies and External Systems Integration"]:
#                             if skill in row and str(row[skill]).strip() and not percentage_pattern.match(str(row[skill])):
#                                 entry["Salesforce Technical Competencies and External Systems Integration"][skill] = str(row[skill])

#                         # Process behavioral competencies and certifications
#                         for skill in categories["Behavioral & Leadership Competencies and Certifications"]:
#                             if skill in row:
#                                 value = row[skill]
#                                 if value == 1:
#                                     entry["Behavioral & Leadership Competencies and Certifications"][skill] = "Certified"
#                                 elif str(value).strip() and not percentage_pattern.match(str(value)):
#                                     entry["Behavioral & Leadership Competencies and Certifications"][skill] = str(value)

#                         structured_data.append(entry)
#                         role_counter += 1

#             except Exception as sheet_error:
#                 logger.error(f"Error processing sheet {sheet}: {str(sheet_error)}")
#                 continue

#         return structured_data[:62]  # Limit to 62 roles
#     except Exception as e:
#         logger.error(f"Error processing Excel file: {str(e)}")
#         raise FileProcessingError("Failed to process Excel file", str(e))

# # Update the extract_skill_matrix function to use the new implementation
# def extract_skill_matrix(file_path, output_json_path):
#     try:
#         # Use context manager for Excel file
#         with pd.ExcelFile(file_path) as xls:
#             skill_matrix_data = extract_clean_skill_matrix(Path(file_path))
        
#         # Save JSON to file
#         with open(output_json_path, "w", encoding="utf-8") as json_file:
#             json.dump(skill_matrix_data, json_file, indent=4, ensure_ascii=False)
        
#         return skill_matrix_data
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"Error processing skill matrix file: {str(e)}")


# def save_as_json(data, output_file):
#     with open(output_file, "w", encoding="utf-8") as json_file:
#         json.dump(data, json_file, indent=4, ensure_ascii=False)


# def generate_document(data, document_type="cover letter"):
#     try:
#         client = Mistral(api_key=MISTRAL_API_KEY)
        
#         prompt = f"""
#         {COVER_LETTER_TEMPLATE}
#         Here is the extracted data:
#         {json.dumps(data, indent=2)}
#         Generate a {document_type} based on the above data.
#         """
        
#         response = client.chat.complete(
#             model=MISTRAL_MODEL,
#             messages=[{"role": "user", "content": prompt}]
#         )
        
#         return response.choices[0].message.content
#     except Exception as e:
#         logger.error(f"Error generating document: {str(e)}")
#         raise HTTPException(status_code=500, detail=f"Error generating document: {str(e)}")
    


# @app.post("/generate-all-documents")
# async def generate_all_documents(data: dict):
#     try:
#         documents = {}
#         timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
#         # Generate both resume and cover letter
#         for doc_type in ["resume", "cover letter"]:
#             content = generate_document(data, doc_type)
#             filename = f"Generated_{doc_type.replace(' ', '_')}_{timestamp}"
            
#             # Save as DOCX
#             docx_path = OUTPUT_DIR / f"{filename}.docx"
#             doc = Document()
#             doc.add_paragraph(content)
#             doc.save(docx_path)
            
#             # Save as PDF
#             pdf_path = OUTPUT_DIR / f"{filename}.pdf"
#             pdf_doc = fitz.open()
#             page = pdf_doc.new_page()
#             page.insert_text((50, 50), content)
#             pdf_doc.save(pdf_path)
#             pdf_doc.close()
            
#             documents[doc_type] = {
#                 "docx": docx_path,
#                 "pdf": pdf_path
#             }
        
#         # Create a zip file containing all documents
#         zip_path = OUTPUT_DIR / f"Generated_Documents_{timestamp}.zip"
#         with zipfile.ZipFile(zip_path, 'w') as zipf:
#             for doc_type, paths in documents.items():
#                 zipf.write(paths["docx"], arcname=f"Generated_{doc_type.replace(' ', '_')}.docx")
#                 zipf.write(paths["pdf"], arcname=f"Generated_{doc_type.replace(' ', '_')}.pdf")
        
#         # Clean up individual files
#         for paths in documents.values():
#             paths["docx"].unlink()
#             paths["pdf"].unlink()
        
#         return FileResponse(
#             path=zip_path,
#             filename=f"Generated_Documents_{timestamp}.zip",
#             media_type="application/zip"
#         )
    
#     except Exception as e:
#         logger.error(f"Error in batch document generation: {str(e)}")
#         raise HTTPException(status_code=500, detail=str(e))

# @app.post("/upload/resume")
# async def upload_resume(file: UploadFile = File(...)):
#     try:
#         # Save uploaded file
#         file_path = UPLOAD_DIR / file.filename
#         with file_path.open("wb") as buffer:
#             shutil.copyfileobj(file.file, buffer)
        
#         # Process file based on type
#         if file_path.suffix.lower() == '.docx':
#             extracted_data = extract_docx(str(file_path))
#         elif file_path.suffix.lower() == '.pdf':
#             extracted_data = extract_pdf(str(file_path))
#         else:
#             raise HTTPException(status_code=400, detail="Unsupported file format")
        
#         # Save extracted data as JSON
#         json_path = UPLOAD_DIR / f"{file_path.stem}.json"
#         with json_path.open("w", encoding="utf-8") as json_file:
#             json.dump(extracted_data, json_file, indent=4, ensure_ascii=False)
        
#         # Return JSON data
#         return extracted_data
    
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))
#     finally:
#         # Clean up uploaded file
#         if file_path.exists():
#             file_path.unlink()

# @app.post("/upload/skillmatrix")
# async def upload_skill_matrix(file: UploadFile = File(...)):
#     try:
#         # Save uploaded file
#         file_path = UPLOAD_DIR / file.filename
#         with file_path.open("wb") as buffer:
#             shutil.copyfileobj(file.file, buffer)
        
#         # Process skill matrix file
#         if file_path.suffix.lower() in ['.xlsx', '.xls']:
#             output_json_path = UPLOAD_DIR / f"{file_path.stem}_skill_matrix.json"
#             extracted_data = extract_skill_matrix(str(file_path), str(output_json_path))
#         else:
#             raise HTTPException(status_code=400, detail="Unsupported file format. Only .xlsx and .xls files are allowed.")
        
#         # Return JSON data
#         return extracted_data
    
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))
#     finally:
#         # Clean up uploaded file
#         if file_path.exists():
#             file_path.unlink()

# @app.post("/combine-data")
# async def combine_data(resume_file: UploadFile = File(...), skill_matrix_file: UploadFile = File(...)):
#     resume_path = None
#     skill_matrix_path = None
    
#     try:
#         # Save uploaded resume file
#         resume_path = UPLOAD_DIR / resume_file.filename
#         with resume_path.open("wb") as buffer:
#             shutil.copyfileobj(resume_file.file, buffer)
        
#         # Save uploaded skill matrix file
#         skill_matrix_path = UPLOAD_DIR / skill_matrix_file.filename
#         with skill_matrix_path.open("wb") as buffer:
#             shutil.copyfileobj(skill_matrix_file.file, buffer)
        
#         # Extract resume data
#         if resume_path.suffix.lower() == '.docx':
#             resume_data = extract_docx(str(resume_path))
#             name, role = extract_name_and_role_from_docx(str(resume_path))
#         elif resume_path.suffix.lower() == '.pdf':
#             resume_data = extract_pdf(str(resume_path))
#             name, role = extract_name_and_role_from_pdf(str(resume_path))
#         else:
#             raise HTTPException(status_code=400, detail="Unsupported resume file format")
        
#         # Extract skill matrix data
#         if skill_matrix_path.suffix.lower() in ['.xlsx', '.xls']:
#             output_json_path = UPLOAD_DIR / f"{skill_matrix_path.stem}_skill_matrix.json"
            
#             # Use a context manager for pandas operations
#             try:
#                 skill_matrix_data = extract_skill_matrix(str(skill_matrix_path), str(output_json_path))
#             except Exception as e:
#                 raise HTTPException(status_code=500, detail=f"Error processing skill matrix: {str(e)}")
#         else:
#             raise HTTPException(status_code=400, detail="Unsupported skill matrix file format")
        
#         # Filter skill matrix data
#         filtered_skill_data = filter_skill_matrix(str(output_json_path), name, role)
        
#         # Combine resume and skill matrix data
#         combined_data = combine_resume_and_skills(resume_data, filtered_skill_data)
        
#         # Save combined data to JSON
#         final_output_path = UPLOAD_DIR / "final_combined_data.json"
#         save_as_json(combined_data, str(final_output_path))
        
#         return combined_data
    
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))
    
#     finally:
#         # Clean up files in a safe way
#         try:
#             if resume_path and resume_path.exists():
#                 resume_path.unlink(missing_ok=True)
#             if skill_matrix_path and skill_matrix_path.exists():
#                 # Close any pandas ExcelFile objects that might be open
#                 pd.Excel._engines.clear()
#                 skill_matrix_path.unlink(missing_ok=True)
#         except Exception as cleanup_error:
#             logger.error(f"Error during cleanup: {cleanup_error}")


# @app.post("/generate-document/{document_type}")
# async def generate_document_endpoint(document_type: str, data: dict):
#     try:
#         # Generate the document content
#         generated_content = generate_document(data, document_type)
        
#         # Create a unique filename
#         timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
#         filename = f"Generated_{document_type.replace(' ', '_')}_{timestamp}"
        
#         # Save as text file first
#         txt_path = OUTPUT_DIR / f"{filename}.txt"
#         with open(txt_path, "w", encoding="utf-8") as f:
#             f.write(generated_content)
        
#         # Convert to PDF using PyMuPDF
#         pdf_path = OUTPUT_DIR / f"{filename}.pdf"
#         doc = fitz.open()
#         page = doc.new_page()
#         page.insert_text((50, 50), generated_content)
#         doc.save(pdf_path)
#         doc.close()
        
#         return FileResponse(
#             path=pdf_path,
#             filename=f"{filename}.pdf",
#             media_type="application/pdf"
#         )
    
#     except Exception as e:
#         logger.error(f"Error in document generation endpoint: {str(e)}")
#         raise HTTPException(status_code=500, detail=str(e))



# if __name__ == "__main__":
#     import uvicorn
#     uvicorn.run(app, host="0.0.0.0", port=8000)




from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import pandas as pd
import re
import json
from pathlib import Path
import shutil
import fitz
from docx import Document
import logging
from mistralai import Mistral
import os
from datetime import datetime
import zipfile

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI()

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Create directories if they don't exist
UPLOAD_DIR = Path("uploads")
OUTPUT_DIR = Path("generated_documents")
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

# Mistral API configuration
MISTRAL_API_KEY = "Vl1QhEbiSbwGhfiXdKxnyz5CYqL1FgC9"
MISTRAL_MODEL = "mistral-large-latest"

# Template instructions
RESUME_TEMPLATE = """
**PROFESSIONAL SUMMARY:**
- Highlight key technical and behavioral competencies with scores above 4.
- Include top 3 certifications to showcase qualifications.
**EDUCATION:**
- List degrees, institutions, and graduation years in reverse chronological order.
**SKILLS:**
- Clearly categorized technical, analytical, and soft skills.
**PROJECT DETAILS:**
Format projects in tabular structure:
**<Project Name>**
| Description | Overview of purpose and outcomes. |
| Role | Role played in the project (e.g., Developer, Team Lead). |
| Technology | Technologies and tools used. |
| Contributions | Specific contributions and responsibilities. |
**TONE & FORMATTING:**
- Maintain a formal, professional, and structured format.
"""

COVER_LETTER_TEMPLATE = """
**COVER LETTER:**
- Start with a professional greeting.
- Introduce yourself and mention the position applied for.
- Highlight key relevant skills and achievements.
- Explain why you're interested in the role and how it aligns with the company's goals.
- End with a call to action and gratitude.
**TONE & FORMATTING:**
- Keep it formal, professional, and concise.
- Ensure proper spacing and clear paragraph structure.
"""

class FileProcessingError(Exception):
    def __init__(self, message, details):
        self.message = message
        self.details = details
        super().__init__(self.message)

def extract_docx(file_path):
    doc = Document(file_path)
    structured_text = {"paragraphs": [], "tables": []}
    
    for para in doc.paragraphs:
        structured_text["paragraphs"].append(para.text.strip())
    
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        structured_text["tables"].append(table_data)
    
    return structured_text

def extract_pdf(file_path):
    doc = fitz.open(file_path)
    structured_text = {"pages": []}
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text")
        structured_text["pages"].append({
            "page_number": page_num + 1,
            "content": text.strip()
        })
    
    return structured_text

def extract_name_and_role_from_docx(file_path):
    doc_data = extract_docx(file_path)
    name, role = None, None

    for para in doc_data["paragraphs"]:
        if not name and any(word in para.lower() for word in ["name", "full name"]):
            name = para.split(":")[-1].strip()
        if not role and any(word in para.lower() for word in ["role", "position", "job title"]):
            role = para.split(":")[-1].strip()

    return name, role

def extract_name_and_role_from_pdf(file_path):
    pdf_data = extract_pdf(file_path)
    name, role = None, None

    for page in pdf_data["pages"]:
        content = page["content"]
        if not name and "name" in content.lower():
            name_line = [line for line in content.split("\n") if "name" in line.lower()][0]
            name = name_line.split(":")[-1].strip()
        if not role and any(word in content.lower() for word in ["role", "position", "job title"]):
            role_line = [line for line in content.split("\n") if any(word in line.lower() for word in ["role", "position", "job title"])][0]
            role = role_line.split(":")[-1].strip()

    return name, role

def extract_clean_skill_matrix(file_path: Path) -> list:
    try:
        xls = pd.ExcelFile(file_path)
        structured_data = []
        role_counter = 1

        percentage_pattern = re.compile(r"\d+%")

        for sheet in xls.sheet_names:
            try:
                df = xls.parse(sheet)
                df = df.dropna(how='all').fillna("")

                columns = list(df.columns)
                if len(columns) < 4:
                    logger.warning(f"Sheet {sheet} has insufficient columns")
                    continue

                first_name_col = columns[0]
                last_name_col = "Unnamed: 1"
                experience_col = columns[2]
                expertise_col = columns[3]

                categories = {
                    "Salesforce Technical Competencies and External Systems Integration": [],
                    "Behavioral & Leadership Competencies and Certifications": []
                }

                category_flag = None
                for col in columns[4:]:
                    if isinstance(col, str):
                        if "%" in col or "Current Capability Score" in col or "Expertise(Years)" in col:
                            continue
                        if "Salesforce Technical Competencies" in col or "External Systems Integration" in col:
                            category_flag = "Salesforce Technical Competencies and External Systems Integration"
                        elif "Behavioral & Leadership Competencies" in col or "SF Certification" in col:
                            category_flag = "Behavioral & Leadership Competencies and Certifications"
                        elif category_flag:
                            categories[category_flag].append(col)

                for _, row in df.iterrows():
                    first_name = str(row[first_name_col]).strip()
                    last_name = str(row[last_name_col]).strip()
                    
                    if first_name and last_name and first_name.lower() != "nan" and last_name.lower() != "nan":
                        entry = {
                            "ID": f"Role_{role_counter}",
                            "Sheet Name": sheet,
                            "First Name": first_name,
                            "Last Name": last_name,
                            "Experience": str(row[experience_col]),
                            "Expertise": str(row[expertise_col]),
                            "Salesforce Technical Competencies and External Systems Integration": {},
                            "Behavioral & Leadership Competencies and Certifications": {}
                        }

                        # Process technical competencies
                        for skill in categories["Salesforce Technical Competencies and External Systems Integration"]:
                            if skill in row and str(row[skill]).strip() and not percentage_pattern.match(str(row[skill])):
                                entry["Salesforce Technical Competencies and External Systems Integration"][skill] = str(row[skill])

                        # Process behavioral competencies and certifications
                        for skill in categories["Behavioral & Leadership Competencies and Certifications"]:
                            if skill in row:
                                value = row[skill]
                                if value == 1:
                                    entry["Behavioral & Leadership Competencies and Certifications"][skill] = "Certified"
                                elif str(value).strip() and not percentage_pattern.match(str(value)):
                                    entry["Behavioral & Leadership Competencies and Certifications"][skill] = str(value)

                        structured_data.append(entry)
                        role_counter += 1

            except Exception as sheet_error:
                logger.error(f"Error processing sheet {sheet}: {str(sheet_error)}")
                continue

        return structured_data[:62]
    except Exception as e:
        logger.error(f"Error processing Excel file: {str(e)}")
        raise FileProcessingError("Failed to process Excel file", str(e))

def extract_skill_matrix(file_path, output_json_path):
    try:
        with pd.ExcelFile(file_path) as xls:
            skill_matrix_data = extract_clean_skill_matrix(Path(file_path))
        
        with open(output_json_path, "w", encoding="utf-8") as json_file:
            json.dump(skill_matrix_data, json_file, indent=4, ensure_ascii=False)
        
        return skill_matrix_data
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing skill matrix file: {str(e)}")

def filter_skill_matrix(skill_json_path, extracted_name, extracted_role):
    with open(skill_json_path, 'r', encoding='utf-8') as file:
        skill_data = json.load(file)

    filtered_by_role = [
        entry for entry in skill_data
        if extracted_role and extracted_role.lower() in entry.get("ID", "").lower()
        or extracted_role.lower() in entry.get("Sheet Name", "").lower()
    ]

    filtered_by_name = []
    for entry in filtered_by_role:
        first_name = entry.get("First Name", "")
        last_name = entry.get("Last Name", "")
        
        if isinstance(first_name, str) and isinstance(last_name, str):
            full_name = f"{first_name} {last_name}".strip().lower()
            if extracted_name and extracted_name.lower() in full_name:
                filtered_by_name.append(entry)

    return filtered_by_name

def combine_resume_and_skills(resume_data, skill_matrix_data):
    combined_data = {
        "resume": resume_data,
        "filtered_skill_matrix": skill_matrix_data
    }
    return combined_data

def save_as_json(data, output_file):
    with open(output_file, "w", encoding="utf-8") as json_file:
        json.dump(data, json_file, indent=4, ensure_ascii=False)

def generate_document(data, document_type="cover letter"):
    try:
        client = Mistral(api_key=MISTRAL_API_KEY)
        
        template = RESUME_TEMPLATE if document_type.lower() == "resume" else COVER_LETTER_TEMPLATE
        
        prompt = f"""
        {template}
        Here is the extracted data:
        {json.dumps(data, indent=2)}
        Generate a {document_type} based on the above data.
        """
        
        response = client.chat.complete(
            model=MISTRAL_MODEL,
            messages=[{"role": "user", "content": prompt}]
        )
        
        return response.choices[0].message.content
    except Exception as e:
        logger.error(f"Error generating document: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating document: {str(e)}")

@app.post("/upload/resume")
async def upload_resume(file: UploadFile = File(...)):
    try:
        file_path = UPLOAD_DIR / file.filename
        with file_path.open("wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        if file_path.suffix.lower() == '.docx':
            extracted_data = extract_docx(str(file_path))
        elif file_path.suffix.lower() == '.pdf':
            extracted_data = extract_pdf(str(file_path))
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format")
        
        json_path = UPLOAD_DIR / f"{file_path.stem}.json"
        with json_path.open("w", encoding="utf-8") as json_file:
            json.dump(extracted_data, json_file, indent=4, ensure_ascii=False)
        
        return extracted_data
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if file_path.exists():
            file_path.unlink()

@app.post("/upload/skillmatrix")
async def upload_skill_matrix(file: UploadFile = File(...)):
    try:
        file_path = UPLOAD_DIR / file.filename
        with file_path.open("wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        if file_path.suffix.lower() in ['.xlsx', '.xls']:
            output_json_path = UPLOAD_DIR / f"{file_path.stem}_skill_matrix.json"
            extracted_data = extract_skill_matrix(str(file_path), str(output_json_path))
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format. Only .xlsx and .xls files are allowed.")
        
        return extracted_data
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if file_path.exists():
            file_path.unlink()

@app.post("/combine-data")
async def combine_data(resume_file: UploadFile = File(...), skill_matrix_file: UploadFile = File(...)):
    resume_path = None
    skill_matrix_path = None
    
    try:
        resume_path = UPLOAD_DIR / resume_file.filename
        with resume_path.open("wb") as buffer:
            shutil.copyfileobj(resume_file.file, buffer)
        
        skill_matrix_path = UPLOAD_DIR / skill_matrix_file.filename
        with skill_matrix_path.open("wb") as buffer:
            shutil.copyfileobj(skill_matrix_file.file, buffer)
        
        if resume_path.suffix.lower() == '.docx':
            resume_data = extract_docx(str(resume_path))
            name, role = extract_name_and_role_from_docx(str(resume_path))
        elif resume_path.suffix.lower() == '.pdf':
            resume_data = extract_pdf(str(resume_path))
            name, role = extract_name_and_role_from_pdf(str(resume_path))
        else:
            raise HTTPException(status_code=400, detail="Unsupported resume file format")
        
        if skill_matrix_path.suffix.lower() in ['.xlsx', '.xls']:
            output_json_path = UPLOAD_DIR / f"{skill_matrix_path.stem}_skill_matrix.json"
            try:
                skill_matrix_data = extract_skill_matrix(str(skill_matrix_path), str(output_json_path))
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Error processing skill matrix: {str(e)}")
        else:
            raise HTTPException(status_code=400, detail="Unsupported skill matrix file format")
        
        # filtered_skill_data = filter_skill_matrix(str(output_json_path), name,
        filtered_skill_data = filter_skill_matrix(str(output_json_path), name, role)
        
        combined_data = combine_resume_and_skills(resume_data, filtered_skill_data)
        
        final_output_path = UPLOAD_DIR / "final_combined_data.json"
        save_as_json(combined_data, str(final_output_path))
        
        return combined_data
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
    finally:
        try:
            if resume_path and resume_path.exists():
                resume_path.unlink(missing_ok=True)
            if skill_matrix_path and skill_matrix_path.exists():
                pd.Excel._engines.clear()
                skill_matrix_path.unlink(missing_ok=True)
        except Exception as cleanup_error:
            logger.error(f"Error during cleanup: {cleanup_error}")

@app.post("/generate-document/{document_type}")
async def generate_document_endpoint(document_type: str, data: dict):
    try:
        generated_content = generate_document(data, document_type)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Generated_{document_type.replace(' ', '_')}_{timestamp}"
        
        docx_path = OUTPUT_DIR / f"{filename}.docx"
        pdf_path = OUTPUT_DIR / f"{filename}.pdf"
        
        # Save as DOCX
        doc = Document()
        doc.add_paragraph(generated_content)
        doc.save(docx_path)
        
        # Convert to PDF
        pdf_doc = fitz.open()
        page = pdf_doc.new_page()
        page.insert_text((50, 50), generated_content)
        pdf_doc.save(pdf_path)
        pdf_doc.close()
        
        # Create zip file
        zip_path = OUTPUT_DIR / f"{filename}.zip"
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            zipf.write(docx_path, arcname=f"{filename}.docx")
            zipf.write(pdf_path, arcname=f"{filename}.pdf")
        
        # Clean up individual files
        docx_path.unlink()
        pdf_path.unlink()
        
        return FileResponse(
            path=zip_path,
            filename=f"{filename}.zip",
            media_type="application/zip"
        )
    
    except Exception as e:
        logger.error(f"Error in document generation endpoint: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate-all-documents")
async def generate_all_documents(data: dict):
    try:
        documents = {}
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Generate both resume and cover letter
        for doc_type in ["resume", "cover letter"]:
            content = generate_document(data, doc_type)
            filename = f"Generated_{doc_type.replace(' ', '_')}_{timestamp}"
            
            # Save as DOCX
            docx_path = OUTPUT_DIR / f"{filename}.docx"
            doc = Document()
            doc.add_paragraph(content)
            doc.save(docx_path)
            
            # Save as PDF
            pdf_path = OUTPUT_DIR / f"{filename}.pdf"
            pdf_doc = fitz.open()
            page = pdf_doc.new_page()
            page.insert_text((50, 50), content)
            pdf_doc.save(pdf_path)
            pdf_doc.close()
            
            documents[doc_type] = {
                "docx": docx_path,
                "pdf": pdf_path
            }
        
        # Create zip file with all documents
        zip_path = OUTPUT_DIR / f"Generated_Documents_{timestamp}.zip"
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for doc_type, paths in documents.items():
                zipf.write(paths["docx"], arcname=f"Generated_{doc_type.replace(' ', '_')}.docx")
                zipf.write(paths["pdf"], arcname=f"Generated_{doc_type.replace(' ', '_')}.pdf")
        
        # Clean up individual files
        for paths in documents.values():
            paths["docx"].unlink()
            paths["pdf"].unlink()
        
        return FileResponse(
            path=zip_path,
            filename=f"Generated_Documents_{timestamp}.zip",
            media_type="application/zip"
        )
    
    except Exception as e:
        logger.error(f"Error in batch document generation: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)