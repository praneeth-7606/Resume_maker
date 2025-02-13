

from fastapi import HTTPException
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks
from fastapi.responses import JSONResponse
import uvicorn
# from fastapi import UploadFile
import fitz  # PyMuPDF for PDF files
from docx import Document  # For DOCX files
import io
import pandas as pd
import json
import io
import re
import fitz  # PyMuPDF
from fpdf import FPDF
import requests
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware  # Import CORSMiddleware

app = FastAPI()

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # Allow requests from this origin
    allow_credentials=True,
    allow_methods=["*"],  # Allow all methods (GET, POST, etc.)
    allow_headers=["*"],  # Allow all headers
)

# ---------------- Global Variables ----------------
candidate_list = []
selected_candidate = {}     # Candidate chosen by the user
resume_files = {}           # Mapping from candidate ID to generated PDF filename
progress_log = {}           # Mapping from candidate ID to progress messages

# ---------------- Utility Functions ----------------

def extract_skill_matrix_from_upload(file: UploadFile):
    """
    Reads an uploaded Excel file (skill matrix) and extracts structured candidate data.
    Raises an HTTPException if the file is empty or not a valid Excel file.
    Returns a list of candidate dictionaries.
    """
    file_bytes = file.file.read()
    if not file_bytes:
        raise HTTPException(
            status_code=400, detail="No file uploaded. Please select a valid Excel (.xlsx) file.")
    try:
        xls = pd.ExcelFile(io.BytesIO(file_bytes), engine="openpyxl")
    except Exception as e:
        raise HTTPException(
            status_code=400, detail="Uploaded file is not a valid Excel (.xlsx) file. Please upload a valid Excel file.")

    structured_data = []
    role_counter = 1
    unique_individuals = set()
    percentage_pattern = re.compile(r"^\d+(\.\d+)?%?$")

    for sheet in xls.sheet_names:
        try:
            df = xls.parse(sheet)
            df = df.dropna(how="all").fillna("")
            columns = list(df.columns)
            if len(columns) < 2:
                continue
            first_name_col = columns[0]
            last_name_col = "Unnamed: 1" if "Unnamed: 1" in columns else columns[1]
            experience_col = columns[2] if len(columns) > 2 else None
            expertise_col = columns[3] if len(columns) > 3 else None

            # Define two merged categories
            categories = {
                "Salesforce Technical Competencies and External Systems Integration": [],
                "Behavioral & Leadership Competencies and Certifications": []
            }
            category_flag = None
            for col in columns[4:]:
                if "%" in col or "Current Capability Score" in col or "Expertise(Years)" in col:
                    continue
                if "Salesforce Technical Competencies" in col or "External Systems Integration" in col:
                    category_flag = "Salesforce Technical Competencies and External Systems Integration"
                elif "Behavioral & Leadership Competencies" in col or "SF Certification" in col:
                    category_flag = "Behavioral & Leadership Competencies and Certifications"
                elif category_flag:
                    categories[category_flag].append(col)

            for _, row in df.iterrows():
                first_name = str(row[first_name_col]).strip()
                last_name = str(row[last_name_col]).strip()
                experience = row[experience_col] if experience_col and isinstance(
                    row[experience_col], (int, float, str)) else ""
                expertise = row[expertise_col] if expertise_col and isinstance(
                    row[expertise_col], (int, float, str)) else ""
                if percentage_pattern.match(first_name) or percentage_pattern.match(last_name):
                    continue
                has_skills = any(
                    isinstance(row[col], (int, float)) and row[col] > 0
                    for col in (categories["Salesforce Technical Competencies and External Systems Integration"] +
                                categories["Behavioral & Leadership Competencies and Certifications"])
                )
                if first_name and last_name and has_skills:
                    full_name = f"{first_name} {last_name}"
                    if full_name not in unique_individuals:
                        unique_individuals.add(full_name)
                        entry = {
                            "ID": f"Role_{role_counter}",
                            "Sheet Name": sheet,
                            "First Name": first_name,
                            "Last Name": last_name,
                            "Experience": experience,
                            "Expertise": expertise,
                            "Salesforce Technical Competencies and External Systems Integration": {
                                skill: row[skill] for skill in categories["Salesforce Technical Competencies and External Systems Integration"]
                                if skill in row and isinstance(row[skill], (int, float)) and row[skill] > 0
                            },
                            "Behavioral & Leadership Competencies and Certifications": {
                                skill: row[skill] for skill in categories["Behavioral & Leadership Competencies and Certifications"]
                                if skill in row and isinstance(row[skill], (int, float)) and row[skill] > 0
                            }
                        }
                        for cert in categories["Behavioral & Leadership Competencies and Certifications"]:
                            if cert in row and row[cert] == 1:
                                entry["Behavioral & Leadership Competencies and Certifications"][cert] = "Certified"
                        structured_data.append(entry)
                        role_counter += 1
                    if len(structured_data) >= 62:
                        break
        except Exception as e:
            print(f"Error processing sheet {sheet}: {e}")
    return structured_data


# def extract_pdf_text_from_upload(file: UploadFile):
#     """
#     Extracts text from an uploaded PDF file using PyMuPDF.
#     """
#     file_bytes = file.file.read()
#     doc = fitz.open(stream=file_bytes, filetype="pdf")
#     text = ""
#     for page in doc:
#         text += page.get_text("text")
#     return text




def extract_pdf_text_from_upload(file: UploadFile):
    """
    Extracts text from an uploaded file (PDF, DOCX, or TXT) using appropriate libraries.
    """
    file_bytes = file.file.read()
    file_extension = file.filename.split(".")[-1].lower()  # Get file extension

    if file_extension == "pdf":
        # Handle PDF files using PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text("text")
        return text

    elif file_extension == "docx":
        # Handle DOCX files using python-docx
        doc = Document(io.BytesIO(file_bytes))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text

    elif file_extension == "txt":
        # Handle TXT files
        return file_bytes.decode("utf-8")

    else:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file format. Please upload a PDF, DOCX, or TXT file."
        )

# ---------------- LLM & PDF Generation Functions ----------------

def generate_new_resume(comp_matrix, resume, new_resume_format):
    """
    Calls an LLM API (via OpenRouter using ASUS TUF API) to reformat the old resume plus candidate skill data
    into a new structured JSON following the provided new resume format.
    """
    system_prompt = """
    You are an AI assistant that reformats resumes into a structured JSON format.
    Take the competency matrix and old resume as input, extract relevant details, and map them to the new resume format.
    Only respond with the new resume format as a JSON object that adheres strictly to the provided JSON Schema. Do not include any extra messages.
    """
    user_prompt = f"""
    Competency Matrix:
    {comp_matrix}
    
    Resume:
    {resume}
    
    Convert this into the following structured format:
    {new_resume_format}
    """
    json_schema = {
        "name": "new_resume",
        "strict": True,
        "schema": {
            "type": "object",
            "properties": {
                "first_name": {"type": "string", "description": "Candidate's first name"},
                "last_name": {"type": "string", "description": "Candidate's last name"},
                "role": {"type": "string", "description": "Candidate's role or position"},
                "professional_summary": {"type": "string", "description": "Professional summary of the candidate"},
                "education": {"type": "string", "description": "Education details of the candidate"},
                "skills": {"type": "string", "description": "Comma separated skills"},
                "projects": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "name": {"type": "string", "description": "Name of the project"},
                            "description": {"type": "string", "description": "Description of the project"},
                            "role": {"type": "string", "description": "Role in the project"},
                            "technology": {"type": "string", "description": "Technologies used in the project"},
                            "role_played": {"type": "string", "description": "Detailed role played in the project"}
                        },
                        "required": ["name", "description", "role", "technology", "role_played"],
                        "additionalProperties": False
                    }
                }
            },
            "required": ["first_name", "last_name", "role", "professional_summary", "education", "skills", "projects"],
            "additionalProperties": False
        }
    }
    payload = {
        "model": "mistralai/mistral-small-24b-instruct-2501",
        "temperature": 0,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        "response_format": {"type": "json_schema", "json_schema": json_schema}
    }
    response = requests.post(
        url="https://openrouter.ai/api/v1/chat/completions",
        headers={
            "Authorization": "Bearer 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",
            "Content-Type": "application/json"
        },
        data=json.dumps(payload)
    )
    response = response.json()
    return response["choices"][0]["message"]["content"]


def generate_cover_letter(resume_data):
    """
    Generates a cover letter by calling an LLM API based on the provided resume data.
    """
    prompt = (
        f"Generate a professional cover letter for a job application at Value Health company sending it to the HR of the company take the recipeinet name from the skill matrix and city as coimbatore  using the following resume details:\n\n"
        f"Name: {resume_data.get('first_name', '')} {resume_data.get('last_name', '')}\n"
        f"Role: {resume_data.get('role', '')}\n"
        f"Professional Summary: {resume_data.get('professional_summary', '')}\n"
        f"Education: {resume_data.get('education', '')}\n"
        f"Skills: {resume_data.get('skills', '')}\n\n"
        f"Cover Letter:"
    )
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": "Bearer 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",
        "Content-Type": "application/json"
    }
    data_payload = {
        "model": "mistralai/mistral-small-24b-instruct-2501",
        "temperature": 0.7,
        "messages": [
            {"role": "system", "content": "You are an AI assistant that generates a professional cover letter."},
            {"role": "user", "content": prompt}
        ]
    }
    response = requests.post(url, headers=headers,
                             data=json.dumps(data_payload))
    response_json = response.json()
    cover_letter_text = response_json["choices"][0]["message"]["content"]
    return cover_letter_text


def generate_resume(data_json, logo_path=None):
    """
    Generates a resume PDF (using FPDF) from the new resume JSON data.
    It also calls the cover letter generation function and appends it on a new page.
    Returns the filename of the generated PDF.
    """
    data = json.loads(data_json)
    pdf = FPDF()
    pdf.add_page()
    # Register Unicode fonts (ensure the .ttf files are accessible)
    pdf.add_font("DejaVu", "", "DejaVuSans.ttf", uni=True)
    pdf.add_font("DejaVu", "B", "DejaVuSans-Bold.ttf", uni=True)
    if logo_path:
        x = 210 - 30 - 10
        y = 10
        try:
            pdf.image(logo_path, x=x, y=y, w=30)
            pdf.set_y(y + 35)
        except RuntimeError:
            pass
    pdf.set_font("DejaVu", "B", 16)
    full_name = f"{data['first_name']} {data['last_name']}"
    pdf.cell(0, 10, full_name, ln=True, align='C')
    pdf.set_font("DejaVu", "B", 14)
    pdf.cell(0, 10, data['role'], ln=True, align='C')
    sections = [
        ('PROFESSIONAL SUMMARY', [data['professional_summary']]),
        ('EDUCATION', [data['education']]),
        ('SKILLS', [data['skills']])
    ]
    for title, content in sections:
        pdf.ln(10)
        pdf.set_font("DejaVu", "B", 12)
        pdf.cell(0, 10, title, ln=True)
        pdf.set_font("DejaVu", "", 12)
        pdf.multi_cell(0, 5, "\n".join(content))
    pdf.ln(10)
    pdf.set_font("DejaVu", "B", 12)
    pdf.cell(0, 10, "PROJECT DETAILS", ln=True)
    pdf.set_font("DejaVu", "", 12)
    for project in data['projects']:
        pdf.ln(5)
        pdf.set_font("DejaVu", "B", 12)
        pdf.cell(0, 10, project['name'], ln=True)
        pdf.set_font("DejaVu", "", 12)
        pdf.multi_cell(0, 5, project['description'])
        pdf.cell(0, 5, f"Role: {project['role']}", ln=True)
        pdf.cell(0, 5, f"Technology: {project['technology']}", ln=True)
        pdf.cell(0, 5, f"Role Played: {project['role_played']}", ln=True)
        pdf.ln(5)
    print("Generating Cover Letter...")
    cover_letter_text = generate_cover_letter(data)
    pdf.add_page()
    pdf.set_font("DejaVu", "B", 16)
    pdf.cell(0, 10, "COVER LETTER", ln=True, align='C')
    pdf.ln(10)
    pdf.set_font("DejaVu", "", 12)
    pdf.multi_cell(0, 5, cover_letter_text)
    filename = f"{data['first_name']}_{data['last_name']}_Resume.pdf"
    pdf.output(filename)
    return filename

# ---------------- Background Task ----------------

def generate_resume_background(comp_matrix, old_resume_text, new_resume_format, candidate_id):
    """
    Background task that:
      1. Updates the progress log at each step.
      2. Calls the LLM to generate the new resume JSON.
      3. Generates the PDF.
      4. Stores the filename in resume_files.
    """
    progress_log[candidate_id] = []
    progress_log[candidate_id].append("Starting resume generation...")

    progress_log[candidate_id].append("Calling LLM to reformat resume...")
    formatted_resume = generate_new_resume(
        comp_matrix, old_resume_text, new_resume_format)
    progress_log[candidate_id].append("LLM reformatting complete.")

    progress_log[candidate_id].append(
        "Generating PDF resume with cover letter...")
    pdf_filename = generate_resume(formatted_resume, logo_path="Logo.png")
    progress_log[candidate_id].append("PDF generation complete.")

    resume_files[candidate_id] = pdf_filename

# ---------------- FastAPI Endpoints ----------------

@app.post("/upload-skill-matrix")
async def upload_skill_matrix(file: UploadFile = File(...)):
    """
    Endpoint to upload the skill matrix file and extract candidate data.
    Returns a JSON response with the list of candidates.
    """
    if not file.filename:
        raise HTTPException(
            status_code=400, detail="No file selected. Please upload a valid Excel (.xlsx) file.")
    contents = await file.read()
    if not contents:
        raise HTTPException(
            status_code=400, detail="Uploaded file is empty. Please upload a valid Excel (.xlsx) file.")
    file_like = io.BytesIO(contents)
    new_file = UploadFile(filename=file.filename, file=file_like)

    global candidate_list
    candidate_list = extract_skill_matrix_from_upload(new_file)
    return JSONResponse(content={"candidates": candidate_list})


@app.post("/select-candidate")
async def select_candidate(candidate_id: str = Form(...)):
    """
    Endpoint to select a candidate from the list.
    Returns a JSON response with the selected candidate's details.
    """
    global candidate_list, selected_candidate
    for candidate in candidate_list:
        if str(candidate.get("ID", "")) == candidate_id:
            selected_candidate = candidate
            break
    return JSONResponse(content={"selected_candidate": selected_candidate})


@app.post("/upload-old-resume")
async def upload_old_resume(file: UploadFile = File(...), background_tasks: BackgroundTasks = None):
    """
    Endpoint to upload the old resume PDF and start the resume generation process.
    Returns a JSON response with the candidate ID and a message indicating the process has started.
    """
    old_resume_text = extract_pdf_text_from_upload(file)
    comp_matrix = json.dumps(selected_candidate)

    new_resume_format = """{
        "first_name": "Firstname here",
        "last_name": "Last Name Here",
        "role": "Role Here",
        "Craft a compelling professional summary based on the skill matrix, emphasizing key technical and behavioral competencies where scores exceed 4. Highlight expertise, experience, and strengths concisely in 50 words, ensuring clarity and impact. Conclude with the  certifications[0:No, 1:Yes], showcasing qualifications that add value to his profile.Dont add any other extra information "


",
        "education": "Education Degree with course Details here...",
        "skills": "Skills here...(Separate it with comma) in bullet points, Get the skills from old resume",
        "projects": [
            {
                "name": "Project1",
                "description": "Developed a web application for...",
                "role": "Lead Developer",
                "technology": "Python, Django, PostgreSQL",
                "role_played": "Designed architecture and led the team."
            },
            {
                "name": "Project2",
                "description": "Mobile app for task management.",
                "role": "Full Stack Developer",
                "technology": "React Native, Node.js",
                "role_played": "Implemented frontend and backend features."
            }
        ],
        certifications: "Add the 3 certifications here"

    }"""
 
    candidate_id = selected_candidate.get("ID", "unknown")
    background_tasks.add_task(generate_resume_background, comp_matrix,
                              old_resume_text, new_resume_format, candidate_id)
    return JSONResponse(content={"candidate_id": candidate_id, "message": "Resume generation started."})


@app.get("/progress")
async def progress(candidate_id: str):
    """
    Polling endpoint: returns JSON with status ("pending" or "ready"),
    the generated PDF filename if ready, and a list of progress messages.
    """
    if candidate_id in resume_files:
        return {"status": "ready", "filename": resume_files[candidate_id], "messages": progress_log.get(candidate_id, [])}
    else:
        return {"status": "pending", "messages": progress_log.get(candidate_id, ["Processing..."])}


@app.get("/download/{pdf_filename}", response_class=FileResponse)
async def download_pdf(pdf_filename: str):
    """
    Serves the generated PDF file for download.
    """
    return FileResponse(pdf_filename, media_type="application/pdf", filename=pdf_filename)

# ---------------- Main ----------------
if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)

# from fastapi import FastAPI, File, UploadFile, HTTPException
# from fastapi.middleware.cors import CORSMiddleware
# from fastapi.responses import FileResponse
# import pandas as pd
# import re
# import json
# from pathlib import Path
# import shutil
# import fitz
# from docx import Document
# import logging
# from mistralai import Mistral
# import os
# from datetime import datetime


# # Set up logging
# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger(__name__)

# app = FastAPI()

# # Configure CORS
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["http://localhost:3000"],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# # Create directories if they don't exist
# UPLOAD_DIR = Path("uploads")
# OUTPUT_DIR = Path("generated_documents")
# UPLOAD_DIR.mkdir(exist_ok=True)
# OUTPUT_DIR.mkdir(exist_ok=True)

# # Mistral API configuration
# MISTRAL_API_KEY = "Vl1QhEbiSbwGhfiXdKxnyz5CYqL1FgC9"
# MISTRAL_MODEL = "mistral-large-latest"

# # # Template instructions
# # COVER_LETTER_TEMPLATE = """
# # ---
# # *COVER LETTER:*  
# # - Start with a *professional greeting* addressing the hiring manager or company.  
# # - Introduce yourself with a *brief background* and mention the *position* you're applying for.  
# # - Highlight *key skills and experiences* relevant to the job, focusing on achievements and specific contributions.  
# # - Explain *why you're interested* in the position and how your goals align with the company's mission.  
# # - Conclude with a *call to action*, expressing eagerness for an interview and thanking them for their time.  
# # ---
# # *TONE & FORMATTING:*  
# # - Maintain a *formal* and *professional tone* throughout.  
# # - Keep it *concise* (preferably within one page) and *focused on relevant experiences*.  
# # - Ensure *proper spacing* and *clear paragraph structure* for easy readability.
# # ---
# # """

# COVER_LETTER_TEMPLATE = """
# **PROFESSIONAL SUMMARY:**
# - Highlight key technical and behavioral competencies with scores above 4.
# - Include top 3 certifications to showcase qualifications.

# **EDUCATION:**
# - List degrees, institutions, and graduation years in reverse chronological order.

# **SKILLS:**
# - Clearly categorized technical, analytical, and soft skills.

# **PROJECT DETAILS:**
# Format projects in tabular structure:

# **<Project Name>**
# | Description | Overview of purpose and outcomes. |
# | Role | Role played in the project (e.g., Developer, Team Lead). |
# | Technology | Technologies and tools used. |
# | Contributions | Specific contributions and responsibilities. |

# **TONE & FORMATTING:**
# - Maintain a formal, professional, and structured format.
# """
# INFERENCE_PARAMS = {
#     "temperature": 0.1,
#     "max_tokens": 1024,
#     "top_p": 0.9,
#     "frequency_penalty": 0,
#     "presence_penalty": 0
# }



# class FileProcessingError(Exception):
#     def __init__(self, message, details):
#         self.message = message
#         self.details = details
#         super().__init__(self.message)

# # app = FastAPI()
# # app.add_middleware(
# #     CORSMiddleware,
# #     allow_origins=["http://localhost:3000"],  # React app URL
# #     allow_credentials=True,
# #     allow_methods=["*"],
# #     allow_headers=["*"],
# # )

# # Create uploads directory if it doesn't exist
# # UPLOAD_DIR = Path("uploads")
# # UPLOAD_DIR.mkdir(exist_ok=True)

# def extract_docx(file_path):
#     doc = Document(file_path)
#     structured_text = {"paragraphs": [], "tables": []}
    
#     for para in doc.paragraphs:
#         structured_text["paragraphs"].append(para.text.strip())
    
#     for table in doc.tables:
#         table_data = []
#         for row in table.rows:
#             row_data = [cell.text.strip() for cell in row.cells]
#             table_data.append(row_data)
#         structured_text["tables"].append(table_data)
    
#     return structured_text

# def extract_pdf(file_path):
#     doc = fitz.open(file_path)
#     structured_text = {"pages": []}
    
#     for page_num in range(len(doc)):
#         page = doc[page_num]
#         text = page.get_text("text")
#         structured_text["pages"].append({
#             "page_number": page_num + 1,
#             "content": text.strip()
#         })
    
#     return structured_text

# def extract_name_and_role_from_docx(file_path):
#     doc_data = extract_docx(file_path)
#     name, role = None, None

#     for para in doc_data["paragraphs"]:
#         if not name and any(word in para.lower() for word in ["name", "full name"]):
#             name = para.split(":")[-1].strip()
#         if not role and any(word in para.lower() for word in ["role", "position", "job title"]):
#             role = para.split(":")[-1].strip()

#     return name, role

# def extract_name_and_role_from_pdf(file_path):
#     pdf_data = extract_pdf(file_path)
#     name, role = None, None

#     for page in pdf_data["pages"]:
#         content = page["content"]
#         if not name and "name" in content.lower():
#             name_line = [line for line in content.split("\n") if "name" in line.lower()][0]
#             name = name_line.split(":")[-1].strip()
#         if not role and any(word in content.lower() for word in ["role", "position", "job title"]):
#             role_line = [line for line in content.split("\n") if any(word in line.lower() for word in ["role", "position", "job title"])][0]
#             role = role_line.split(":")[-1].strip()

#     return name, role

# def filter_skill_matrix(skill_json_path, extracted_name, extracted_role):
#     with open(skill_json_path, 'r', encoding='utf-8') as file:
#         skill_data = json.load(file)

#     # Filter by role
#     filtered_by_role = [
#         entry for entry in skill_data
#         if extracted_role and extracted_role.lower() in entry.get("ID", "").lower()
#         or extracted_role.lower() in entry.get("Sheet Name", "").lower()
#     ]

#     # Further filter by name
#     filtered_by_name = []
#     for entry in filtered_by_role:
#         first_name = entry.get("First Name", "")
#         last_name = entry.get("Last Name", "")
        
#         if isinstance(first_name, str) and isinstance(last_name, str):
#             full_name = f"{first_name} {last_name}".strip().lower()
#             if extracted_name and extracted_name.lower() in full_name:
#                 filtered_by_name.append(entry)

#     return filtered_by_name

# def combine_resume_and_skills(resume_data, skill_matrix_data):
#     combined_data = {
#         "resume": resume_data,
#         "filtered_skill_matrix": skill_matrix_data
#     }
#     return combined_data
# def extract_clean_skill_matrix(file_path: Path) -> list:
#     try:
#         xls = pd.ExcelFile(file_path)
#         structured_data = []
#         role_counter = 1

#         percentage_pattern = re.compile(r"\d+%")

#         for sheet in xls.sheet_names:
#             try:
#                 df = xls.parse(sheet)
#                 df = df.dropna(how='all').fillna("")

#                 columns = list(df.columns)
#                 if len(columns) < 4:
#                     logger.warning(f"Sheet {sheet} has insufficient columns")
#                     continue

#                 first_name_col = columns[0]
#                 last_name_col = "Unnamed: 1"
#                 experience_col = columns[2]
#                 expertise_col = columns[3]

#                 categories = {
#                     "Salesforce Technical Competencies and External Systems Integration": [],
#                     "Behavioral & Leadership Competencies and Certifications": []
#                 }

#                 category_flag = None
#                 for col in columns[4:]:
#                     if isinstance(col, str):
#                         if "%" in col or "Current Capability Score" in col or "Expertise(Years)" in col:
#                             continue
#                         if "Salesforce Technical Competencies" in col or "External Systems Integration" in col:
#                             category_flag = "Salesforce Technical Competencies and External Systems Integration"
#                         elif "Behavioral & Leadership Competencies" in col or "SF Certification" in col:
#                             category_flag = "Behavioral & Leadership Competencies and Certifications"
#                         elif category_flag:
#                             categories[category_flag].append(col)

#                 for _, row in df.iterrows():
#                     first_name = str(row[first_name_col]).strip()
#                     last_name = str(row[last_name_col]).strip()
                    
#                     if first_name and last_name and first_name.lower() != "nan" and last_name.lower() != "nan":
#                         entry = {
#                             "ID": f"Role_{role_counter}",
#                             "Sheet Name": sheet,
#                             "First Name": first_name,
#                             "Last Name": last_name,
#                             "Experience": str(row[experience_col]),
#                             "Expertise": str(row[expertise_col]),
#                             "Salesforce Technical Competencies and External Systems Integration": {},
#                             "Behavioral & Leadership Competencies and Certifications": {}
#                         }

#                         # Process technical competencies
#                         for skill in categories["Salesforce Technical Competencies and External Systems Integration"]:
#                             if skill in row and str(row[skill]).strip() and not percentage_pattern.match(str(row[skill])):
#                                 entry["Salesforce Technical Competencies and External Systems Integration"][skill] = str(row[skill])

#                         # Process behavioral competencies and certifications
#                         for skill in categories["Behavioral & Leadership Competencies and Certifications"]:
#                             if skill in row:
#                                 value = row[skill]
#                                 if value == 1:
#                                     entry["Behavioral & Leadership Competencies and Certifications"][skill] = "Certified"
#                                 elif str(value).strip() and not percentage_pattern.match(str(value)):
#                                     entry["Behavioral & Leadership Competencies and Certifications"][skill] = str(value)

#                         structured_data.append(entry)
#                         role_counter += 1

#             except Exception as sheet_error:
#                 logger.error(f"Error processing sheet {sheet}: {str(sheet_error)}")
#                 continue

#         return structured_data[:62]  # Limit to 62 roles
#     except Exception as e:
#         logger.error(f"Error processing Excel file: {str(e)}")
#         raise FileProcessingError("Failed to process Excel file", str(e))

# # Update the extract_skill_matrix function to use the new implementation
# def extract_skill_matrix(file_path, output_json_path):
#     try:
#         # Use context manager for Excel file
#         with pd.ExcelFile(file_path) as xls:
#             skill_matrix_data = extract_clean_skill_matrix(Path(file_path))
        
#         # Save JSON to file
#         with open(output_json_path, "w", encoding="utf-8") as json_file:
#             json.dump(skill_matrix_data, json_file, indent=4, ensure_ascii=False)
        
#         return skill_matrix_data
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"Error processing skill matrix file: {str(e)}")


# def save_as_json(data, output_file):
#     with open(output_file, "w", encoding="utf-8") as json_file:
#         json.dump(data, json_file, indent=4, ensure_ascii=False)

# def generate_document(data, document_type="cover letter"):
#     try:
#         client = Mistral(api_key=MISTRAL_API_KEY)
        
#         prompt = f"""
#         {COVER_LETTER_TEMPLATE}
#         Here is the extracted data:
#         {json.dumps(data, indent=2)}
#         Generate a {document_type} based on the above data.
#         """
        
#         response = client.chat.complete(
#             model=MISTRAL_MODEL,
#             messages=[{"role": "user", "content": prompt}],
#         **INFERENCE_PARAMS
#         )
        
#         return response.choices[0].message.content
#     except Exception as e:
#         logger.error(f"Error generating document: {str(e)}")
#         raise HTTPException(status_code=500, detail=f"Error generating document: {str(e)}")
    
# @app.post("/upload/resume")
# async def upload_resume(file: UploadFile = File(...)):
#     try:
#         # Save uploaded file
#         file_path = UPLOAD_DIR / file.filename
#         with file_path.open("wb") as buffer:
#             shutil.copyfileobj(file.file, buffer)
        
#         # Process file based on type
#         if file_path.suffix.lower() == '.docx':
#             extracted_data = extract_docx(str(file_path))
#         elif file_path.suffix.lower() == '.pdf':
#             extracted_data = extract_pdf(str(file_path))
#         else:
#             raise HTTPException(status_code=400, detail="Unsupported file format")
        
#         # Save extracted data as JSON
#         json_path = UPLOAD_DIR / f"{file_path.stem}.json"
#         with json_path.open("w", encoding="utf-8") as json_file:
#             json.dump(extracted_data, json_file, indent=4, ensure_ascii=False)
        
#         # Return JSON data
#         return extracted_data
    
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))
#     finally:
#         # Clean up uploaded file
#         if file_path.exists():
#             file_path.unlink()

# @app.post("/upload/skillmatrix")
# async def upload_skill_matrix(file: UploadFile = File(...)):
#     try:
#         # Save uploaded file
#         file_path = UPLOAD_DIR / file.filename
#         with file_path.open("wb") as buffer:
#             shutil.copyfileobj(file.file, buffer)
        
#         # Process skill matrix file
#         if file_path.suffix.lower() in ['.xlsx', '.xls']:
#             output_json_path = UPLOAD_DIR / f"{file_path.stem}_skill_matrix.json"
#             extracted_data = extract_skill_matrix(str(file_path), str(output_json_path))
#         else:
#             raise HTTPException(status_code=400, detail="Unsupported file format. Only .xlsx and .xls files are allowed.")
        
#         # Return JSON data
#         return extracted_data
    
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))
#     finally:
#         # Clean up uploaded file
#         if file_path.exists():
#             file_path.unlink()

# @app.post("/combine-data")
# async def combine_data(resume_file: UploadFile = File(...), skill_matrix_file: UploadFile = File(...)):
#     resume_path = None
#     skill_matrix_path = None
    
#     try:
#         # Save uploaded resume file
#         resume_path = UPLOAD_DIR / resume_file.filename
#         with resume_path.open("wb") as buffer:
#             shutil.copyfileobj(resume_file.file, buffer)
        
#         # Save uploaded skill matrix file
#         skill_matrix_path = UPLOAD_DIR / skill_matrix_file.filename
#         with skill_matrix_path.open("wb") as buffer:
#             shutil.copyfileobj(skill_matrix_file.file, buffer)
        
#         # Extract resume data
#         if resume_path.suffix.lower() == '.docx':
#             resume_data = extract_docx(str(resume_path))
#             name, role = extract_name_and_role_from_docx(str(resume_path))
#         elif resume_path.suffix.lower() == '.pdf':
#             resume_data = extract_pdf(str(resume_path))
#             name, role = extract_name_and_role_from_pdf(str(resume_path))
#         else:
#             raise HTTPException(status_code=400, detail="Unsupported resume file format")
        
#         # Extract skill matrix data
#         if skill_matrix_path.suffix.lower() in ['.xlsx', '.xls']:
#             output_json_path = UPLOAD_DIR / f"{skill_matrix_path.stem}_skill_matrix.json"
            
#             # Use a context manager for pandas operations
#             try:
#                 skill_matrix_data = extract_skill_matrix(str(skill_matrix_path), str(output_json_path))
#             except Exception as e:
#                 raise HTTPException(status_code=500, detail=f"Error processing skill matrix: {str(e)}")
#         else:
#             raise HTTPException(status_code=400, detail="Unsupported skill matrix file format")
        
#         # Filter skill matrix data
#         filtered_skill_data = filter_skill_matrix(str(output_json_path), name, role)
        
#         # Combine resume and skill matrix data
#         combined_data = combine_resume_and_skills(resume_data, filtered_skill_data)
        
#         # Save combined data to JSON
#         final_output_path = UPLOAD_DIR / "final_combined_data.json"
#         save_as_json(combined_data, str(final_output_path))
        
#         return combined_data
    
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))
    
#     finally:
#         # Clean up files in a safe way
#         try:
#             if resume_path and resume_path.exists():
#                 resume_path.unlink(missing_ok=True)
#             if skill_matrix_path and skill_matrix_path.exists():
#                 # Close any pandas ExcelFile objects that might be open
#                 pd.Excel._engines.clear()
#                 skill_matrix_path.unlink(missing_ok=True)
#         except Exception as cleanup_error:
#             logger.error(f"Error during cleanup: {cleanup_error}")


# @app.post("/generate-document/{document_type}")
# async def generate_document_endpoint(document_type: str, data: dict):
#     try:
#         # Generate the document content
#         generated_content = generate_document(data, document_type)
        
#         # Create a unique filename
#         timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
#         filename = f"Generated_{document_type.replace(' ', '_')}_{timestamp}"
        
#         # Save as text file first
#         txt_path = OUTPUT_DIR / f"{filename}.txt"
#         with open(txt_path, "w", encoding="utf-8") as f:
#             f.write(generated_content)
        
#         # Convert to PDF using PyMuPDF
#         pdf_path = OUTPUT_DIR / f"{filename}.pdf"
#         doc = fitz.open()
#         page = doc.new_page()
#         page.insert_text((50, 50), generated_content)
#         doc.save(pdf_path)
#         doc.close()
        
#         return FileResponse(
#             path=pdf_path,
#             filename=f"{filename}.pdf",
#             media_type="application/pdf"
#         )
    
#     except Exception as e:
#         logger.error(f"Error in document generation endpoint: {str(e)}")
#         raise HTTPException(status_code=500, detail=str(e))



# if __name__ == "__main__":
#     import uvicorn
#     uvicorn.run(app, host="0.0.0.0", port=8000)
